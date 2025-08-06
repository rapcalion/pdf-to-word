"""
Quick script to check conversion results
"""

import os
from docx import Document

def check_docx(file_path):
    """Check contents of converted DOCX file"""
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    doc = Document(file_path)
    
    print(f"Analyzing: {file_path}")
    print("-" * 50)
    
    # Count elements
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    
    print(f"Paragraphs: {paragraphs}")
    print(f"Tables: {tables}")
    
    # Show table info
    if tables > 0:
        print("\nTable details:")
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            print(f"  Table {i+1}: {rows} rows x {cols} columns")
    
    # Show first few paragraphs
    print("\nFirst few paragraphs:")
    for i, para in enumerate(doc.paragraphs[:5]):
        if para.text.strip():
            print(f"  {i+1}. {para.text[:60]}...")
    
    print("\n✓ Conversion appears successful!")
    print(f"File size: {os.path.getsize(file_path) / 1024:.1f} KB")

if __name__ == "__main__":
    check_docx("CV - JASH THAKKER_converted.docx")
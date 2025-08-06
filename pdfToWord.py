import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
import os


def pdf_to_word(pdf_path, output_path):
    """
    Convert PDF to Word document with improved formatting and table detection
    """
    # Open PDF
    pdf_document = fitz.open(pdf_path)
    
    # Create new Word document
    doc = Document()
    
    # Process each page
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        
        # Add page break for pages after the first
        if page_num > 0:
            doc.add_page_break()
        
        # Try to find tables first
        tables = page.find_tables()
        table_areas = []
        
        # Process tables
        for table in tables:
            table_areas.append(table.bbox)
            process_table(doc, table)
        
        # Extract text blocks with formatting
        blocks = page.get_text("dict")
        
        # Group blocks by y-coordinate for better layout
        sorted_blocks = sorted(blocks["blocks"], key=lambda b: b.get("bbox", [0, 0, 0, 0])[1])
        
        for block in sorted_blocks:
            # Skip text that's already in tables
            if is_in_table_area(block, table_areas):
                continue
                
            if "lines" in block:  # Text block
                process_text_block(doc, block)
            elif "image" in block:  # Image block
                process_image_block(doc, block, page)
    
    # Save the document
    doc.save(output_path)
    pdf_document.close()
    print(f"Conversion completed: {output_path}")


def process_table(doc, table):
    """
    Extract and format tables
    """
    try:
        # Extract table data
        table_data = table.extract()
        
        if not table_data:
            return
        
        # Create Word table
        word_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        word_table.style = 'Table Grid'
        
        # Fill table data
        for row_idx, row in enumerate(table_data):
            for col_idx, cell_text in enumerate(row):
                if cell_text:
                    cell = word_table.cell(row_idx, col_idx)
                    cell.text = str(cell_text).strip()
        
        # Add spacing after table
        doc.add_paragraph("")
        
    except Exception as e:
        print(f"Error processing table: {e}")


def is_in_table_area(block, table_areas):
    """
    Check if a text block is within any table area
    """
    if not table_areas or "bbox" not in block:
        return False
    
    block_bbox = block["bbox"]
    
    for table_bbox in table_areas:
        # Check if block overlaps with table area
        if (block_bbox[0] < table_bbox[2] and block_bbox[2] > table_bbox[0] and
            block_bbox[1] < table_bbox[3] and block_bbox[3] > table_bbox[1]):
            return True
    
    return False


def process_text_block(doc, block):
    """
    Process text blocks with improved formatting and spacing
    """
    # Group lines by similar y-coordinates to detect paragraphs
    lines = block.get("lines", [])
    if not lines:
        return
    
    # Process each line
    for line_idx, line in enumerate(lines):
        # Create paragraph for each line
        paragraph = doc.add_paragraph()
        
        # Check if this looks like a heading (larger font, bold)
        is_heading = False
        line_text = ""
        
        for span in line["spans"]:
            text = span["text"]
            line_text += text
            
            if text.strip():
                run = paragraph.add_run(text)
                
                # Apply formatting
                font = run.font
                font.name = span.get("font", "Arial")
                font_size = span.get("size", 12)
                font.size = Pt(font_size)
                
                # Check for heading characteristics
                if font_size > 14:
                    is_heading = True
                
                # Handle font flags
                flags = span.get("flags", 0)
                if flags & 2**4:  # Bold
                    font.bold = True
                    is_heading = True
                if flags & 2**1:  # Italic
                    font.italic = True
                if flags & 2**2:  # Underline
                    font.underline = True
        
        # Apply heading style if detected
        if is_heading and line_text.strip():
            if len(line_text.strip()) < 100:  # Likely a heading
                paragraph.style = 'Heading 2'
        
        # Add extra spacing for paragraphs
        if line_idx == len(lines) - 1:  # Last line of block
            doc.add_paragraph("")


def process_image_block(doc, block, page):
    """
    Extract and add images to Word document
    """
    try:
        # Get image data
        img_index = block["image"]
        base_image = page.parent.extract_image(img_index)
        image_bytes = base_image["image"]
        
        # Save temporary image
        temp_img_path = f"temp_image_{img_index}.png"
        with open(temp_img_path, "wb") as img_file:
            img_file.write(image_bytes)
        
        # Add image to document
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(temp_img_path, width=Inches(6))
        
        # Clean up temporary file
        os.remove(temp_img_path)
        
    except Exception as e:
        print(f"Error processing image: {e}")


def main():
    # Example usage
    pdf_path = input("Enter PDF file path: ").strip()
    
    if not os.path.exists(pdf_path):
        print("File not found!")
        return
    
    # Generate output filename
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_path = f"{base_name}_converted.docx"
    
    try:
        pdf_to_word(pdf_path, output_path)
    except Exception as e:
        print(f"Error during conversion: {e}")


if __name__ == "__main__":
    main()
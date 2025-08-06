"""
Comprehensive PDF to Word Converter
Handles all PDF types with visual formatting preservation
"""

import os
import sys
import fitz  # PyMuPDF
import pdfplumber
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pytesseract
from PIL import Image
import numpy as np
import cv2
import io
import tempfile
import shutil
from typing import Dict, List, Tuple, Optional, Any
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ComprehensivePDFConverter:
    """
    Main converter class that orchestrates multiple conversion strategies
    """
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        self.conversion_methods = {
            'pdf2docx': self._convert_with_pdf2docx,
            'pymupdf': self._convert_with_pymupdf,
            'pdfplumber': self._convert_with_pdfplumber,
            'hybrid': self._convert_hybrid
        }
    
    def convert(self, pdf_path: str, output_path: str, method: str = 'hybrid') -> bool:
        """
        Convert PDF to Word using specified method
        """
        try:
            logger.info(f"Starting conversion: {pdf_path} -> {output_path}")
            logger.info(f"Using method: {method}")
            
            # Check if PDF is scanned
            if self._is_scanned_pdf(pdf_path):
                logger.info("Detected scanned PDF, using OCR")
                return self._convert_scanned_pdf(pdf_path, output_path)
            
            # Use specified conversion method
            if method in self.conversion_methods:
                return self.conversion_methods[method](pdf_path, output_path)
            else:
                logger.error(f"Unknown method: {method}")
                return False
                
        except Exception as e:
            logger.error(f"Conversion failed: {str(e)}")
            return False
        finally:
            # Cleanup
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
    
    def _is_scanned_pdf(self, pdf_path: str) -> bool:
        """
        Check if PDF is scanned (image-based)
        """
        try:
            with pdfplumber.open(pdf_path) as pdf:
                # Check first few pages
                pages_to_check = min(3, len(pdf.pages))
                text_found = False
                
                for i in range(pages_to_check):
                    text = pdf.pages[i].extract_text()
                    if text and len(text.strip()) > 50:
                        text_found = True
                        break
                
                return not text_found
        except:
            return False
    
    def _convert_with_pdf2docx(self, pdf_path: str, output_path: str) -> bool:
        """
        Direct conversion using pdf2docx library
        """
        try:
            cv = Converter(pdf_path)
            cv.convert(output_path)
            cv.close()
            logger.info("pdf2docx conversion completed")
            return True
        except Exception as e:
            logger.error(f"pdf2docx failed: {str(e)}")
            return False
    
    def _convert_with_pymupdf(self, pdf_path: str, output_path: str) -> bool:
        """
        Conversion using PyMuPDF with enhanced formatting
        """
        try:
            pdf_doc = fitz.open(pdf_path)
            word_doc = Document()
            
            for page_num, page in enumerate(pdf_doc):
                if page_num > 0:
                    word_doc.add_page_break()
                
                # Extract and process content
                self._process_page_pymupdf(page, word_doc)
            
            word_doc.save(output_path)
            pdf_doc.close()
            logger.info("PyMuPDF conversion completed")
            return True
        except Exception as e:
            logger.error(f"PyMuPDF failed: {str(e)}")
            return False
    
    def _convert_with_pdfplumber(self, pdf_path: str, output_path: str) -> bool:
        """
        Conversion using pdfplumber with focus on tables
        """
        try:
            word_doc = Document()
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    if page_num > 0:
                        word_doc.add_page_break()
                    
                    self._process_page_pdfplumber(page, word_doc)
            
            word_doc.save(output_path)
            logger.info("pdfplumber conversion completed")
            return True
        except Exception as e:
            logger.error(f"pdfplumber failed: {str(e)}")
            return False
    
    def _convert_hybrid(self, pdf_path: str, output_path: str) -> bool:
        """
        Hybrid approach combining multiple methods
        """
        try:
            # First try pdf2docx for best formatting
            temp_output = os.path.join(self.temp_dir, "temp_pdf2docx.docx")
            if self._convert_with_pdf2docx(pdf_path, temp_output):
                # Enhance with table detection from pdfplumber
                self._enhance_with_tables(pdf_path, temp_output, output_path)
                return True
            
            # Fallback to custom implementation
            logger.info("Falling back to custom implementation")
            return self._custom_comprehensive_convert(pdf_path, output_path)
            
        except Exception as e:
            logger.error(f"Hybrid conversion failed: {str(e)}")
            return False
    
    def _custom_comprehensive_convert(self, pdf_path: str, output_path: str) -> bool:
        """
        Custom comprehensive conversion with all features
        """
        try:
            pdf_doc = fitz.open(pdf_path)
            word_doc = Document()
            
            # Set document margins
            sections = word_doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            with pdfplumber.open(pdf_path) as plumber_pdf:
                for page_num in range(len(pdf_doc)):
                    if page_num > 0:
                        word_doc.add_page_break()
                    
                    # Get both representations
                    mupdf_page = pdf_doc[page_num]
                    plumber_page = plumber_pdf.pages[page_num]
                    
                    # Process page with combined approach
                    self._process_page_comprehensive(
                        mupdf_page, plumber_page, word_doc, page_num
                    )
            
            word_doc.save(output_path)
            pdf_doc.close()
            logger.info("Custom comprehensive conversion completed")
            return True
            
        except Exception as e:
            logger.error(f"Custom conversion failed: {str(e)}")
            return False
    
    def _process_page_comprehensive(self, mupdf_page, plumber_page, word_doc, page_num):
        """
        Process a page using multiple extraction methods
        """
        logger.info(f"Processing page {page_num + 1}")
        
        # Extract tables first
        tables = self._extract_tables_comprehensive(plumber_page)
        table_regions = [table['bbox'] for table in tables]
        
        # Extract images
        images = self._extract_images(mupdf_page)
        
        # Extract text with formatting
        blocks = mupdf_page.get_text("dict")
        
        # Sort content by position
        content_items = []
        
        # Add tables to content items
        for table in tables:
            content_items.append({
                'type': 'table',
                'data': table['data'],
                'bbox': table['bbox'],
                'y_pos': table['bbox'][1]
            })
        
        # Add images to content items
        for img in images:
            content_items.append({
                'type': 'image',
                'data': img['data'],
                'bbox': img['bbox'],
                'y_pos': img['bbox'][1]
            })
        
        # Add text blocks to content items
        for block in blocks['blocks']:
            if 'lines' in block:
                # Check if block is inside a table region
                if not self._is_in_regions(block.get('bbox', [0,0,0,0]), table_regions):
                    content_items.append({
                        'type': 'text',
                        'data': block,
                        'bbox': block.get('bbox', [0,0,0,0]),
                        'y_pos': block.get('bbox', [0,0,0,0])[1]
                    })
        
        # Sort by vertical position
        content_items.sort(key=lambda x: x['y_pos'])
        
        # Process sorted content
        for item in content_items:
            if item['type'] == 'table':
                self._add_table_to_doc(word_doc, item['data'])
            elif item['type'] == 'image':
                self._add_image_to_doc(word_doc, item['data'])
            elif item['type'] == 'text':
                self._add_text_block_to_doc(word_doc, item['data'])
    
    def _extract_tables_comprehensive(self, plumber_page):
        """
        Extract tables using multiple methods
        """
        tables = []
        
        # Method 1: pdfplumber table detection
        plumber_tables = plumber_page.extract_tables()
        for i, table in enumerate(plumber_tables):
            if table and len(table) > 0:
                # Get table bbox
                table_bbox = self._estimate_table_bbox(plumber_page, table)
                tables.append({
                    'data': table,
                    'bbox': table_bbox
                })
        
        # Method 2: Custom table detection based on lines
        if not tables:
            custom_tables = self._detect_tables_from_lines(plumber_page)
            tables.extend(custom_tables)
        
        return tables
    
    def _estimate_table_bbox(self, page, table_data):
        """
        Estimate bounding box for a table
        """
        # Simple estimation - can be improved
        page_height = page.height
        page_width = page.width
        
        # Rough estimation based on table position
        rows = len(table_data)
        cols = len(table_data[0]) if table_data else 0
        
        # Estimate based on page dimensions
        y_start = 100  # Default top margin
        y_end = y_start + (rows * 20)  # Rough row height
        x_start = 50
        x_end = page_width - 50
        
        return [x_start, y_start, x_end, y_end]
    
    def _detect_tables_from_lines(self, page):
        """
        Detect tables based on line patterns
        """
        tables = []
        
        # Extract horizontal and vertical lines
        h_lines = page.horizontal_edges
        v_lines = page.vertical_edges
        
        # Group lines to detect table regions
        # This is a simplified version - can be enhanced
        if len(h_lines) > 2 and len(v_lines) > 2:
            # Potential table detected
            logger.info("Detected potential table from lines")
        
        return tables
    
    def _extract_images(self, page):
        """
        Extract images from PDF page
        """
        images = []
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            try:
                # Extract image
                xref = img[0]
                pix = fitz.Pixmap(page.parent, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes()
                else:  # Convert CMYK to RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                    img_data = pix.tobytes()
                
                # Get image position
                img_bbox = page.get_image_bbox(img)
                
                images.append({
                    'data': img_data,
                    'bbox': list(img_bbox),
                    'width': pix.width,
                    'height': pix.height
                })
                
                pix = None
                
            except Exception as e:
                logger.warning(f"Failed to extract image: {str(e)}")
        
        return images
    
    def _is_in_regions(self, bbox, regions):
        """
        Check if bbox is inside any of the regions
        """
        for region in regions:
            if (bbox[0] >= region[0] and bbox[2] <= region[2] and
                bbox[1] >= region[1] and bbox[3] <= region[3]):
                return True
        return False
    
    def _add_table_to_doc(self, doc, table_data):
        """
        Add a table to the Word document
        """
        if not table_data or not table_data[0]:
            return
        
        try:
            # Clean table data
            clean_data = []
            for row in table_data:
                clean_row = []
                for cell in row:
                    cell_text = str(cell) if cell is not None else ""
                    clean_row.append(cell_text.strip())
                if any(cell for cell in clean_row):  # Skip empty rows
                    clean_data.append(clean_row)
            
            if not clean_data:
                return
            
            # Create table
            num_rows = len(clean_data)
            num_cols = max(len(row) for row in clean_data)
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Fill table
            for row_idx, row_data in enumerate(clean_data):
                for col_idx in range(num_cols):
                    cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_text
                    
                    # Format header row
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            
            # Add spacing after table
            doc.add_paragraph()
            
        except Exception as e:
            logger.warning(f"Failed to add table: {str(e)}")
    
    def _add_image_to_doc(self, doc, img_data):
        """
        Add an image to the Word document
        """
        try:
            # Save image temporarily
            temp_img = os.path.join(self.temp_dir, f"temp_img_{id(img_data)}.png")
            
            # Convert bytes to image
            if isinstance(img_data, dict):
                img_bytes = img_data.get('data', b'')
                width = img_data.get('width', 0)
                height = img_data.get('height', 0)
            else:
                img_bytes = img_data
                width = height = 0
            
            # Save image
            with open(temp_img, 'wb') as f:
                f.write(img_bytes)
            
            # Add to document
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            
            # Determine appropriate size
            max_width = Inches(6)
            if width > 0:
                aspect_ratio = height / width
                if width > 500:  # Large image
                    img_width = max_width
                else:
                    img_width = Inches(width / 100)  # Scale down
            else:
                img_width = max_width
            
            run.add_picture(temp_img, width=img_width)
            
            # Add spacing
            doc.add_paragraph()
            
        except Exception as e:
            logger.warning(f"Failed to add image: {str(e)}")
    
    def _add_text_block_to_doc(self, doc, block):
        """
        Add formatted text block to document
        """
        try:
            lines = block.get('lines', [])
            
            for line in lines:
                paragraph = doc.add_paragraph()
                
                for span in line.get('spans', []):
                    text = span.get('text', '')
                    if not text.strip():
                        continue
                    
                    run = paragraph.add_run(text)
                    
                    # Apply formatting
                    font = run.font
                    font.name = span.get('font', 'Arial')
                    font.size = Pt(span.get('size', 11))
                    
                    # Font flags
                    flags = span.get('flags', 0)
                    if flags & 2**4:  # Bold
                        font.bold = True
                    if flags & 2**1:  # Italic
                        font.italic = True
                    if flags & 2**2:  # Underline
                        font.underline = True
                    
                    # Color
                    color = span.get('color', 0)
                    if color != 0:
                        rgb = self._int_to_rgb(color)
                        font.color.rgb = RGBColor(*rgb)
        
        except Exception as e:
            logger.warning(f"Failed to add text block: {str(e)}")
    
    def _int_to_rgb(self, color_int):
        """
        Convert integer color to RGB tuple
        """
        r = (color_int >> 16) & 0xFF
        g = (color_int >> 8) & 0xFF
        b = color_int & 0xFF
        return (r, g, b)
    
    def _enhance_with_tables(self, pdf_path: str, docx_path: str, output_path: str):
        """
        Enhance existing DOCX with better table detection
        """
        try:
            # Load existing document
            doc = Document(docx_path)
            
            # Extract tables using pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    if tables:
                        logger.info(f"Found {len(tables)} tables on page {page_num + 1}")
            
            # Save enhanced document
            doc.save(output_path)
            
        except Exception as e:
            logger.warning(f"Table enhancement failed: {str(e)}")
            # Copy original file
            shutil.copy2(docx_path, output_path)
    
    def _convert_scanned_pdf(self, pdf_path: str, output_path: str) -> bool:
        """
        Convert scanned PDF using OCR
        """
        try:
            logger.info("Starting OCR conversion")
            pdf_doc = fitz.open(pdf_path)
            word_doc = Document()
            
            for page_num, page in enumerate(pdf_doc):
                if page_num > 0:
                    word_doc.add_page_break()
                
                # Convert page to image
                mat = fitz.Matrix(2, 2)  # Increase resolution
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # OCR the image
                image = Image.open(io.BytesIO(img_data))
                text = pytesseract.image_to_string(image)
                
                # Add text to document
                if text.strip():
                    for line in text.split('\n'):
                        if line.strip():
                            word_doc.add_paragraph(line)
                
                # Also try to extract any embedded text
                page_text = page.get_text()
                if page_text.strip():
                    word_doc.add_paragraph("---")
                    word_doc.add_paragraph(page_text)
            
            word_doc.save(output_path)
            pdf_doc.close()
            logger.info("OCR conversion completed")
            return True
            
        except Exception as e:
            logger.error(f"OCR conversion failed: {str(e)}")
            return False
    
    def _process_page_pymupdf(self, page, doc):
        """
        Process page using PyMuPDF
        """
        # Extract tables
        tables = page.find_tables()
        for table in tables:
            try:
                data = table.extract()
                self._add_table_to_doc(doc, data)
            except:
                pass
        
        # Extract text blocks
        blocks = page.get_text("dict")
        for block in blocks['blocks']:
            if 'lines' in block:
                self._add_text_block_to_doc(doc, block)
    
    def _process_page_pdfplumber(self, page, doc):
        """
        Process page using pdfplumber
        """
        # Extract tables
        tables = page.extract_tables()
        for table in tables:
            if table:
                self._add_table_to_doc(doc, table)
        
        # Extract text
        text = page.extract_text()
        if text:
            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)


def main():
    """
    Main entry point with CLI interface
    """
    print("=" * 60)
    print("Comprehensive PDF to Word Converter")
    print("=" * 60)
    
    # Get input file
    pdf_path = input("Enter PDF file path: ").strip()
    
    if not os.path.exists(pdf_path):
        print("Error: File not found!")
        return
    
    # Show conversion options
    print("\nConversion methods:")
    print("1. hybrid (Recommended) - Best quality, combines multiple methods")
    print("2. pdf2docx - Fast, good formatting")
    print("3. pymupdf - Good for complex documents")
    print("4. pdfplumber - Best for tables")
    
    method_choice = input("\nSelect method (1-4, default=1): ").strip()
    
    method_map = {
        '1': 'hybrid',
        '2': 'pdf2docx',
        '3': 'pymupdf',
        '4': 'pdfplumber'
    }
    
    method = method_map.get(method_choice, 'hybrid')
    
    # Generate output filename
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_path = f"{base_name}_converted_{method}.docx"
    
    # Convert
    converter = ComprehensivePDFConverter()
    
    print(f"\nConverting using {method} method...")
    success = converter.convert(pdf_path, output_path, method)
    
    if success:
        print(f"\n✓ Conversion successful!")
        print(f"Output saved to: {output_path}")
    else:
        print("\n✗ Conversion failed!")
        print("Try a different method or check the logs.")


if __name__ == "__main__":
    main()
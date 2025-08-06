import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import pandas as pd


def pdf_to_word_advanced(pdf_path, output_path):
    """
    Advanced PDF to Word converter with better table and layout detection
    """
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        
        if page_num > 0:
            doc.add_page_break()
        
        print(f"Processing page {page_num + 1}...")
        
        # Method 1: Try PyMuPDF table detection
        tables_found = False
        try:
            tables = page.find_tables()
            if tables:
                tables_found = True
                print(f"Found {len(tables)} tables using PyMuPDF")
                process_page_with_tables(doc, page, tables)
            else:
                # Method 2: Use text analysis for table detection
                detect_tables_from_text(doc, page)
        except Exception as e:
            print(f"Table detection failed: {e}")
            # Fallback to basic text extraction
            extract_text_basic(doc, page)
    
    doc.save(output_path)
    pdf_document.close()
    print(f"Advanced conversion completed: {output_path}")


def process_page_with_tables(doc, page, tables):
    """
    Process page with detected tables
    """
    # Get all text blocks
    blocks = page.get_text("dict")
    table_areas = [table.bbox for table in tables]
    
    # Sort blocks by position
    sorted_blocks = sorted(blocks["blocks"], key=lambda b: (b.get("bbox", [0, 0, 0, 0])[1], b.get("bbox", [0, 0, 0, 0])[0]))
    
    table_index = 0
    
    for block in sorted_blocks:
        # Check if we should insert a table here
        if table_index < len(tables):
            block_y = block.get("bbox", [0, 0, 0, 0])[1]
            table_y = tables[table_index].bbox[1]
            
            if block_y > table_y:
                # Insert table before this block
                create_advanced_table(doc, tables[table_index])
                table_index += 1
        
        # Process text block if not in table area
        if not is_in_any_table_area(block, table_areas):
            process_text_block_advanced(doc, block)
    
    # Insert remaining tables
    while table_index < len(tables):
        create_advanced_table(doc, tables[table_index])
        table_index += 1


def create_advanced_table(doc, table):
    """
    Create a well-formatted table in Word
    """
    try:
        table_data = table.extract()
        if not table_data or not table_data[0]:
            return
        
        # Filter out empty rows and columns
        filtered_data = []
        for row in table_data:
            if any(cell and str(cell).strip() for cell in row):
                filtered_data.append(row)
        
        if not filtered_data:
            return
        
        # Determine number of columns
        max_cols = max(len(row) for row in filtered_data)
        
        # Create Word table
        word_table = doc.add_table(rows=len(filtered_data), cols=max_cols)
        word_table.style = 'Table Grid'
        word_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Fill table data
        for row_idx, row in enumerate(filtered_data):
            for col_idx in range(max_cols):
                cell_text = ""
                if col_idx < len(row) and row[col_idx]:
                    cell_text = str(row[col_idx]).strip()
                
                cell = word_table.cell(row_idx, col_idx)
                cell.text = cell_text
                
                # Format first row as header
                if row_idx == 0 and cell_text:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
        
        # Add spacing
        doc.add_paragraph("")
        
    except Exception as e:
        print(f"Error creating table: {e}")


def detect_tables_from_text(doc, page):
    """
    Detect tables by analyzing text positioning and alignment
    """
    blocks = page.get_text("dict")
    
    # Group text by lines and analyze spacing
    lines_data = []
    
    for block in blocks["blocks"]:
        if "lines" not in block:
            continue
            
        for line in block["lines"]:
            line_text = ""
            line_bbox = None
            
            for span in line["spans"]:
                line_text += span["text"]
                if line_bbox is None:
                    line_bbox = span["bbox"]
                else:
                    # Extend bbox
                    line_bbox = [
                        min(line_bbox[0], span["bbox"][0]),
                        min(line_bbox[1], span["bbox"][1]),
                        max(line_bbox[2], span["bbox"][2]),
                        max(line_bbox[3], span["bbox"][3])
                    ]
            
            if line_text.strip():
                lines_data.append({
                    'text': line_text,
                    'bbox': line_bbox,
                    'y': line_bbox[1] if line_bbox else 0
                })
    
    # Sort by y position
    lines_data.sort(key=lambda x: x['y'])
    
    # Detect potential table areas
    potential_tables = detect_table_patterns(lines_data)
    
    if potential_tables:
        for table_lines in potential_tables:
            create_table_from_lines(doc, table_lines)
    else:
        # No tables detected, process as regular text
        for line_data in lines_data:
            p = doc.add_paragraph(line_data['text'])


def detect_table_patterns(lines_data):
    """
    Analyze text patterns to detect tables
    """
    table_lines = []
    current_table = []
    
    for i, line in enumerate(lines_data):
        text = line['text']
        
        # Check if line looks like table content
        # Criteria: multiple spaces, tabs, or consistent column patterns
        if (text.count('  ') >= 2 or  # Multiple double spaces
            text.count('\t') >= 1 or   # Contains tabs
            len([x for x in text.split() if x]) >= 3):  # Multiple words
            
            current_table.append(line)
        else:
            # End of potential table
            if len(current_table) >= 2:  # Minimum 2 rows for a table
                table_lines.append(current_table)
            current_table = []
    
    # Check last table
    if len(current_table) >= 2:
        table_lines.append(current_table)
    
    return table_lines


def create_table_from_lines(doc, table_lines):
    """
    Create table from detected text lines
    """
    try:
        # Parse each line into columns
        table_data = []
        
        for line in table_lines:
            text = line['text']
            # Split by multiple spaces or tabs
            columns = [col.strip() for col in text.replace('\t', '  ').split('  ') if col.strip()]
            if columns:
                table_data.append(columns)
        
        if not table_data:
            return
        
        # Determine max columns
        max_cols = max(len(row) for row in table_data)
        
        # Create Word table
        word_table = doc.add_table(rows=len(table_data), cols=max_cols)
        word_table.style = 'Table Grid'
        
        # Fill data
        for row_idx, row in enumerate(table_data):
            for col_idx in range(max_cols):
                cell_text = row[col_idx] if col_idx < len(row) else ""
                word_table.cell(row_idx, col_idx).text = cell_text
        
        doc.add_paragraph("")
        
    except Exception as e:
        print(f"Error creating table from lines: {e}")


def process_text_block_advanced(doc, block):
    """
    Advanced text processing with better formatting
    """
    lines = block.get("lines", [])
    if not lines:
        return
    
    for line in lines:
        paragraph = doc.add_paragraph()
        
        for span in line["spans"]:
            text = span["text"]
            if text.strip():
                run = paragraph.add_run(text)
                
                # Apply formatting
                font = run.font
                font.name = span.get("font", "Arial")
                font.size = Pt(span.get("size", 12))
                
                # Handle font flags
                flags = span.get("flags", 0)
                if flags & 2**4:  # Bold
                    font.bold = True
                if flags & 2**1:  # Italic
                    font.italic = True
                if flags & 2**2:  # Underline
                    font.underline = True


def is_in_any_table_area(block, table_areas):
    """
    Check if block is within any table area
    """
    if not table_areas or "bbox" not in block:
        return False
    
    block_bbox = block["bbox"]
    
    for table_bbox in table_areas:
        if (block_bbox[0] < table_bbox[2] and block_bbox[2] > table_bbox[0] and
            block_bbox[1] < table_bbox[3] and block_bbox[3] > table_bbox[1]):
            return True
    
    return False


def extract_text_basic(doc, page):
    """
    Basic text extraction fallback
    """
    text = page.get_text()
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)


def main():
    print("Advanced PDF to Word Converter")
    pdf_path = input("Enter PDF file path: ").strip()
    
    if not os.path.exists(pdf_path):
        print("File not found!")
        return
    
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_path = f"{base_name}_advanced_converted.docx"
    
    try:
        pdf_to_word_advanced(pdf_path, output_path)
    except Exception as e:
        print(f"Error during conversion: {e}")


if __name__ == "__main__":
    main()